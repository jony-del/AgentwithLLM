from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("deliverables/ReAct模式伪代码.docx")

NAVY = RGBColor(31, 78, 121)
BLUE = RGBColor(0, 92, 197)
GREEN = RGBColor(0, 112, 60)
STRING = RGBColor(163, 55, 55)
PURPLE = RGBColor(112, 48, 160)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(100, 108, 118)
LIGHT = "E8EEF5"


CODE = '''class ReActAgent:

    async def run(self, task, history=None,
                  should_cancel=None, deadline=None):
        # 1. 初始化运行环境与上下文
        reset_run_state()
        sandbox.reset()

        messages = [SystemMessage(config.system_prompt)]
        messages += await memory.recall(task)
        messages += await build_project_context()

        if history:
            messages += remove_old_system_context(history)

        user_message = UserMessage(task)
        prompt_hook = await hooks.on_user_prompt(user_message)

        if prompt_hook.blocked:
            return stop(prompt_hook.reason)

        messages.append(prompt_hook.transformed_message)

        # 2. ReAct 主循环
        step = 0
        stop_block_count = 0

        while True:
            # 2.1 安全退出检查
            if should_cancel():
                return stop("用户取消")

            if config.max_steps is not None \\
                    and step >= config.max_steps:
                return stop("超过最大步骤数")

            if deadline_has_passed(deadline):
                return stop("超过运行时间限制")

            if near_deadline(deadline):
                messages.append(
                    SystemMessage("停止调用工具并尽快总结")
                )

            step += 1

            # 2.2 必要时压缩上下文
            if compression.should_compact(messages):
                pre_hook = await hooks.before_compact(messages)

                if not pre_hook.blocked:
                    messages = await compression.compact(
                        messages,
                        preserve_system_context=True,
                        preserve_recent_messages=True,
                        reattach_read_files=True,
                    )
                    await transcript.save_compaction_boundary(messages)
                    await hooks.after_compact(messages)

            # 2.3 Reason：模型分析状态并决定下一步
            try:
                result = await provider.complete(
                    messages=messages,
                    tools=tool_registry.schemas(),
                    model=config.model,
                    stream=config.stream,
                )

            except ContextTooLongError:
                # 溢出恢复：压缩、删除旧轮次、缩短大消息后重试
                messages = await compression.reactive_compact(messages)
                result = await retry_after_shrinking_context(messages)

            assistant_message = AssistantMessage(
                content=result.content,
                tool_calls=result.tool_calls,
                thinking_blocks=result.thinking_blocks,
                provider_state=result.provider_state,
            )
            messages.append(assistant_message)
            await transcript.append(assistant_message)
            await hooks.after_sampling(assistant_message)

            # 3. 无工具调用：模型认为任务可以结束
            if not result.tool_calls:
                stop_hook = await hooks.before_stop(
                    messages=messages,
                    answer=result.content,
                )

                if stop_hook.blocked \\
                        and stop_block_count < config.max_stop_blocks:
                    stop_block_count += 1
                    messages.append(UserMessage(
                        stop_hook.additional_context
                        or "任务尚未完成，请继续工作"
                    ))
                    continue

                await memory.extract_and_save(messages)
                await logger.save_final(result.content)

                return AgentRunResult(
                    answer=result.content,
                    messages=messages,
                    steps=step,
                    run_id=logger.run_id,
                )

            # 4. Act：通过权限与沙箱检查后执行工具
            tool_results = await tool_executor.execute_many(
                result.tool_calls,
                parallel=config.parallel_tools,
                max_workers=config.max_tool_workers,
                should_cancel=should_cancel,
            )

            # 5. Observe：把每个工具结果写回消息历史
            for tool_call, tool_result in zip(
                    result.tool_calls, tool_results):
                observation = ToolMessage(
                    name=tool_result.name,
                    content=tool_result.content,
                    tool_call_id=tool_call.id,
                    metadata={"success": tool_result.ok},
                )
                messages.append(observation)
                await transcript.append(observation)

                record_file_read(tool_call, tool_result)
                await notify_lsp_if_file_changed(
                    tool_call, tool_result
                )

            # 工具执行期间的新用户输入加入下一轮
            messages += drain_midturn_user_messages()

            # 回到循环：Observation 成为下一轮 Reason 的输入
'''


KEYWORDS = {
    "class", "async", "def", "await", "if", "else", "elif", "while",
    "for", "in", "try", "except", "return", "continue", "True", "False",
    "None", "and", "or", "not",
}
TOKEN_RE = re.compile(
    r'(#.*$)|("(?:[^"\\]|\\.)*")|\b(' + "|".join(sorted(KEYWORDS, key=len, reverse=True)) + r')\b'
)


def set_font(run, name: str, size: float, color: RGBColor, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def set_cell_like_padding(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "100")
    ind.set(qn("w:right"), "100")
    ppr.append(ind)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, "Microsoft YaHei", 8, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    set_font(run, "Microsoft YaHei", 8, MUTED)


def add_code_line(doc: Document, number: int, line: str) -> None:
    p = doc.add_paragraph(style="CodeLine")
    shade_paragraph(p, "F6F8FA")
    set_cell_like_padding(p)

    number_run = p.add_run(f"{number:>3}  ")
    set_font(number_run, "Consolas", 7.8, MUTED)

    cursor = 0
    for match in TOKEN_RE.finditer(line):
        if match.start() > cursor:
            run = p.add_run(line[cursor:match.start()])
            set_font(run, "Consolas", 8.4, INK)

        token = match.group(0)
        if match.group(1):
            color = GREEN
        elif match.group(2):
            color = STRING
        else:
            color = BLUE if token not in {"True", "False", "None"} else PURPLE

        run = p.add_run(token)
        set_font(run, "Consolas", 8.4, color, bold=bool(match.group(3)))
        cursor = match.end()

    if cursor < len(line):
        run = p.add_run(line[cursor:])
        set_font(run, "Consolas", 8.4, INK)

    if not line:
        run = p.add_run(" ")
        set_font(run, "Consolas", 8.4, INK)


def build_document() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, size, before, after in (
        ("Heading 1", 14, 8, 4),
        ("Heading 2", 11, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = doc.styles.add_style("CodeLine", WD_STYLE_TYPE.PARAGRAPH)
    code_style.base_style = normal
    code_style.paragraph_format.space_before = Pt(0)
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 0.92
    code_style.paragraph_format.keep_together = True
    code_style.paragraph_format.widow_control = False

    header = section.header.paragraphs[0]
    header.text = "AgentwithLLM  |  ReAct 核心流程"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.runs[0], "Microsoft YaHei", 8, MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("ReAct 模式伪代码")
    set_font(run, "Microsoft YaHei", 20, NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(7)
    run = subtitle.add_run("基于 agent_core/react.py 的核心执行逻辑（打印精简版）")
    set_font(run, "Microsoft YaHei", 9, MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.0
    labels = [
        ("Reason", BLUE, "模型判断下一步"),
        ("Act", PURPLE, "执行工具"),
        ("Observe", GREEN, "结果写回上下文"),
    ]
    for index, (label, color, desc) in enumerate(labels):
        if index:
            sep = p.add_run("   →   ")
            set_font(sep, "Microsoft YaHei", 9, MUTED)
        r = p.add_run(label)
        set_font(r, "Microsoft YaHei", 9, color, bold=True)
        r = p.add_run(f"：{desc}")
        set_font(r, "Microsoft YaHei", 9, INK)

    doc.add_heading("核心伪代码", level=1)
    for index, line in enumerate(CODE.splitlines(), start=1):
        add_code_line(doc, index, line)

    doc.add_heading("循环摘要", level=1)
    summary = [
        ("1", BLUE, "准备上下文", "加载系统提示、记忆、项目信息和历史消息。"),
        ("2", BLUE, "Reason", "LLM 根据当前消息判断：继续调用工具，或直接回答。"),
        ("3", PURPLE, "Act", "工具通过权限和沙箱检查后执行；无资源冲突时可并行。"),
        ("4", GREEN, "Observe", "工具结果写回消息历史，成为下一轮模型输入。"),
        ("5", STRING, "Finish", "没有工具调用时，经过 Stop Hook 后输出最终答案。"),
    ]
    for num, color, label, desc in summary:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}. {label}  ")
        set_font(r, "Microsoft YaHei", 9.2, color, bold=True)
        r = p.add_run(desc)
        set_font(r, "Microsoft YaHei", 9.2, INK)

    note = doc.add_paragraph()
    shade_paragraph(note, LIGHT)
    set_cell_like_padding(note)
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(0)
    r = note.add_run("关键保护：")
    set_font(r, "Microsoft YaHei", 9, NAVY, bold=True)
    r = note.add_run("上下文压缩、取消/超时、最大步骤、权限审批、沙箱、Hooks、会话持久化。")
    set_font(r, "Microsoft YaHei", 9, INK)

    doc.core_properties.title = "ReAct 模式伪代码"
    doc.core_properties.subject = "基于 agent_core/react.py 的核心执行逻辑"
    doc.core_properties.author = "AgentwithLLM"
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
